from io import BytesIO
from html.parser import HTMLParser
from typing import List, Dict

from docx import Document
from docx.shared import RGBColor

from ..models.schemas import ExportDocxRequest


COLOR_INSERT = RGBColor(30, 58, 138)   # blue-800 approx
COLOR_DELETE = RGBColor(185, 28, 28)   # red-700 approx
COLOR_MOVE = RGBColor(6, 95, 70)       # emerald-800 approx


class SpanCollector(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.stack: List[str] = []
        self.chunks: List[Dict[str, object]] = []

    def handle_starttag(self, tag: str, attrs):
        if tag == "span":
            cls = ""
            for k, v in attrs:
                if k == "class":
                    cls = v or ""
                    break
            self.stack.append(cls)
        elif tag == "p":
            # treat paragraph tag as soft container; no-op
            pass

    def handle_endtag(self, tag: str):
        if tag == "span" and self.stack:
            self.stack.pop()

    def handle_data(self, data: str):
        current = self.stack[-1] if self.stack else ""
        if data:
            self.chunks.append({"text": data, "cls": current})


def _apply_run_style(run, cls: str) -> None:
    if not cls:
        return
    if "diff-insert" in cls:
        run.font.color.rgb = COLOR_INSERT
    if "diff-delete" in cls:
        run.font.color.rgb = COLOR_DELETE
        run.font.strike = True
    if "diff-move" in cls:
        run.font.color.rgb = COLOR_MOVE
    if "line-through" in cls and "diff-delete" not in cls:
        run.font.strike = True


def _add_paragraph_from_html(doc: Document, html: str) -> None:
    parser = SpanCollector()
    parser.feed(html)
    p = doc.add_paragraph("")
    for chunk in parser.chunks:
        run = p.add_run(str(chunk["text"]))
        _apply_run_style(run, str(chunk.get("cls", "")))


def render_docx(payload: ExportDocxRequest) -> bytes:
    doc = Document()
    # Header
    heading = doc.add_paragraph()
    heading.add_run("Erik's Super Compare\n").bold = True
    heading.add_run("Redlines so good, even Litera blushes.\n")

    if payload.diffHtmlByParagraph:
        for html in payload.diffHtmlByParagraph:
            _add_paragraph_from_html(doc, html)
    elif payload.diffRunsByParagraph:
        # Future: build from structured runs
        for para in payload.diffRunsByParagraph:
            p = doc.add_paragraph("")
            for run_def in para or []:
                text = str(run_def.get("text", ""))
                cls = str(run_def.get("class", ""))
                run = p.add_run(text)
                _apply_run_style(run, cls)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


