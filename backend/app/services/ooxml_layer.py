"""OOXML layer for native DOCX structure reading and revision tracking.

Provides low-level access to DOCX OOXML structure including sections, paragraphs,
runs, formatting, and revision tracking elements.
"""
import hashlib
import io
from dataclasses import dataclass
from typing import List, Literal, Optional, Tuple, Dict, Any
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as _DocxDocument
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph as _DocxParagraph
from docx.text.run import Run as _DocxRun

from ..utils.text_ops import normalize_for_compare


# OOXML namespace prefixes
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


@dataclass
class RunInfo:
    """Information about a text run within a paragraph."""
    text: str
    rPr_xml: bytes  # Preserved formatting XML as bytes
    start_pos: int  # Character offset in paragraph text
    xml_element: Any  # Reference to the <w:r> element


@dataclass
class ParagraphInfo:
    """Information about a paragraph with its runs and metadata."""
    text: str
    normalized: str  # Normalized text for comparison
    runs: List[RunInfo]
    style_sig: str  # SHA-1 hash of style properties
    xml_path: tuple  # (section_idx, block_idx, para_idx) for tracking
    xml_element: Any  # Reference to the <w:p> element
    metadata: Dict[str, Any]  # Additional metadata (numbering, style, etc.)


@dataclass
class Block:
    """A structural block in the document (section, paragraph, table, header, footer)."""
    kind: Literal["section", "paragraph", "table", "header", "footer"]
    xml_element: Any  # Underlying XML element
    paragraphs: List[ParagraphInfo]
    metadata: Dict[str, Any]  # Style, numbering, position info


def load_docx(data: bytes) -> Document:
    """Load a DOCX document from bytes."""
    return Document(io.BytesIO(data))


def _get_run_properties_xml(run_element: Any) -> bytes:
    """Extract <w:rPr> XML from a run element."""
    rPr = run_element.find(f"./{W_NS}rPr")
    if rPr is not None:
        return ET.tostring(run_element.find(f"./{W_NS}rPr"), encoding="unicode").encode("utf-8")
    return b"<w:rPr/>"


def _compute_style_signature(paragraph_element: Any) -> str:
    """Compute SHA-1 hash of paragraph style properties for comparison."""
    # Extract relevant style properties
    props: List[str] = []
    
    # Paragraph style
    pStyle = paragraph_element.find(f"./{W_NS}pPr/{W_NS}pStyle")
    if pStyle is not None:
        props.append(f"pStyle:{pStyle.get(qn('w:val'), '')}")
    
    # Numbering properties
    numPr = paragraph_element.find(f"./{W_NS}pPr/{W_NS}numPr")
    if numPr is not None:
        ilvl = numPr.find(f"./{W_NS}ilvl")
        numId = numPr.find(f"./{W_NS}numId")
        if ilvl is not None:
            props.append(f"ilvl:{ilvl.get(qn('w:val'), '')}")
        if numId is not None:
            props.append(f"numId:{numId.get(qn('w:val'), '')}")
    
    # Indentation
    ind = paragraph_element.find(f"./{W_NS}pPr/{W_NS}ind")
    if ind is not None:
        left = ind.get(qn('w:left'), '')
        right = ind.get(qn('w:right'), '')
        if left:
            props.append(f"indLeft:{left}")
        if right:
            props.append(f"indRight:{right}")
    
    # Alignment
    jc = paragraph_element.find(f"./{W_NS}pPr/{W_NS}jc")
    if jc is not None:
        props.append(f"jc:{jc.get(qn('w:val'), '')}")
    
    sig_str = "|".join(sorted(props))
    return hashlib.sha1(sig_str.encode("utf-8")).hexdigest()


def enumerate_runs(paragraph: _DocxParagraph) -> List[RunInfo]:
    """Enumerate runs within a paragraph, preserving formatting."""
    runs: List[RunInfo] = []
    current_pos = 0
    
    # Access the underlying XML element
    p_element = paragraph._element
    
    for run_element in p_element.findall(f"./{W_NS}r"):
        # Extract text from all <w:t> elements in this run
        text_parts: List[str] = []
        for t_elem in run_element.findall(f"./{W_NS}t"):
            if t_elem.text:
                text_parts.append(t_elem.text)
        
        run_text = "".join(text_parts)
        
        # Get run properties XML
        rPr_xml = _get_run_properties_xml(run_element)
        
        runs.append(RunInfo(
            text=run_text,
            rPr_xml=rPr_xml,
            start_pos=current_pos,
            xml_element=run_element
        ))
        
        current_pos += len(run_text)
    
    return runs


def enumerate_paragraphs(block: Any, section_idx: int, block_idx: int) -> List[ParagraphInfo]:
    """Enumerate paragraphs within a block (paragraph element, table cell, header, footer)."""
    paragraphs: List[ParagraphInfo] = []
    para_idx = 0
    
    # Handle different block types
    if isinstance(block, _DocxParagraph):
        # Single paragraph block
        p_element = block._element
        runs = enumerate_runs(block)
        text = block.text or ""
        normalized = normalize_for_compare(text)
        style_sig = _compute_style_signature(p_element)
        
        # Extract metadata
        metadata: Dict[str, Any] = {
            "style": block.style.name if block.style else "Normal",
            "alignment": str(block.alignment) if block.alignment else None,
        }
        
        # Check for numbering
        numPr = p_element.find(f"./{W_NS}pPr/{W_NS}numPr")
        if numPr is not None:
            metadata["is_list"] = True
        
        paragraphs.append(ParagraphInfo(
            text=text,
            normalized=normalized,
            runs=runs,
            style_sig=style_sig,
            xml_path=(section_idx, block_idx, para_idx),
            xml_element=p_element,
            metadata=metadata
        ))
    elif isinstance(block, Table):
        # Table: enumerate paragraphs in each cell
        for row_idx, row in enumerate(block.rows):
            for col_idx, cell in enumerate(row.cells):
                for cell_para in cell.paragraphs:
                    p_element = cell_para._element
                    runs = enumerate_runs(cell_para)
                    text = cell_para.text or ""
                    normalized = normalize_for_compare(text)
                    style_sig = _compute_style_signature(p_element)
                    
                    metadata: Dict[str, Any] = {
                        "kind": "table",
                        "row": row_idx,
                        "col": col_idx,
                        "style": cell_para.style.name if cell_para.style else "Normal",
                    }
                    
                    paragraphs.append(ParagraphInfo(
                        text=text,
                        normalized=normalized,
                        runs=runs,
                        style_sig=style_sig,
                        xml_path=(section_idx, block_idx, para_idx),
                        xml_element=p_element,
                        metadata=metadata
                    ))
                    para_idx += 1
    else:
        # Handle other block types (headers, footers)
        # For now, treat as paragraph container
        if hasattr(block, "paragraphs"):
            for para in block.paragraphs:
                p_element = para._element
                runs = enumerate_runs(para)
                text = para.text or ""
                normalized = normalize_for_compare(text)
                style_sig = _compute_style_signature(p_element)
                
                metadata: Dict[str, Any] = {
                    "style": para.style.name if para.style else "Normal",
                }
                
                paragraphs.append(ParagraphInfo(
                    text=text,
                    normalized=normalized,
                    runs=runs,
                    style_sig=style_sig,
                    xml_path=(section_idx, block_idx, para_idx),
                    xml_element=p_element,
                    metadata=metadata
                ))
                para_idx += 1
    
    return paragraphs


def enumerate_blocks(document: Document) -> List[Block]:
    """Enumerate structural blocks in document order (sections, paragraphs, tables)."""
    blocks: List[Block] = []
    section_idx = 0
    
    # Access document body
    body = document.element.body
    
    # Iterate through direct children of body
    block_idx = 0
    for child in body.iterchildren():
        tag = child.tag
        paragraphs: List[ParagraphInfo] = []
        
        if tag.endswith("}p"):
            # Paragraph block
            para = _DocxParagraph(child, document)
            paragraphs = enumerate_paragraphs(para, section_idx, block_idx)
            
            blocks.append(Block(
                kind="paragraph",
                xml_element=child,
                paragraphs=paragraphs,
                metadata={}
            ))
            block_idx += 1
        elif tag.endswith("}tbl"):
            # Table block
            table = Table(child, document)
            paragraphs = enumerate_paragraphs(table, section_idx, block_idx)
            
            blocks.append(Block(
                kind="table",
                xml_element=child,
                paragraphs=paragraphs,
                metadata={}
            ))
            block_idx += 1
        elif tag.endswith("}sectPr"):
            # Section properties - mark section boundary
            # Sections are implicit in DOCX structure
            pass
    
    return blocks


def read_document_xml(document: Document) -> Dict[str, Any]:
    """Extract comprehensive document structure including headers, footers, sections."""
    structure: Dict[str, Any] = {
        "blocks": [],
        "sections": [],
        "headers": {},
        "footers": {},
        "numbering": {},
    }
    
    # Enumerate main body blocks
    blocks = enumerate_blocks(document)
    structure["blocks"] = blocks
    
    # Extract headers and footers
    # python-docx doesn't expose headers/footers easily, so we need to access OOXML directly
    try:
        part = document.part
        if hasattr(part, "related_parts"):
            for rel in part.related_parts.values():
                if "header" in rel.__class__.__name__.lower():
                    # This is a header part
                    pass  # TODO: Extract header paragraphs
                elif "footer" in rel.__class__.__name__.lower():
                    # This is a footer part
                    pass  # TODO: Extract footer paragraphs
    except Exception:
        # Headers/footers extraction is complex; defer for now
        pass
    
    return structure


def write_revisions(document: Document, revisions: List[Dict[str, Any]]) -> bytes:
    """Apply revision tracking elements to document and return as bytes.
    
    Revisions is a list of operations with keys:
    - kind: 'insert', 'delete', 'move_from', 'move_to'
    - xml_path: tuple indicating where to apply
    - element: XML element to wrap/insert
    - author: revision author (default: "Erik's Super Compare")
    - date: revision date (optional)
    """
    # This is a placeholder; actual implementation will be in ooxml_rewriter.py
    # This function is here for API completeness
    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()

